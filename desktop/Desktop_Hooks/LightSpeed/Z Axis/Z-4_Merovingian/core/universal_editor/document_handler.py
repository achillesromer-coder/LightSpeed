"""
Document File Handler - V1.0.0
Handler for document files (PDF, DOCX) - preview only

Author: LightSpeed Team
Date: December 27, 2025
"""

import tkinter as tk
from tkinter import ttk, messagebox
from pathlib import Path
from typing import List, Optional

from .file_handler import FileHandler, EditorCapability, EditorAction


# ==============================================================================
# Document File Handler
# ==============================================================================

class DocumentFileHandler(FileHandler):
    """
    Handler for document files (PDF, DOCX)

    Features:
    - PDF preview (text extraction)
    - DOCX preview (text extraction)
    - Metadata display
    - Page navigation
    - Search within document
    """

    def __init__(self):
        """Initialize document file handler"""
        super().__init__()

        self.file_type_name = "Document"
        self.supported_extensions = ['.pdf', '.docx', '.doc']

        self.capabilities = [
            EditorCapability.READ,
            EditorCapability.PREVIEW,
            EditorCapability.SEARCH,
        ]

        self.current_widget: Optional[tk.Text] = None
        self.current_file: Optional[Path] = None
        self.current_pages: List[str] = []
        self.current_page_index: int = 0

    def can_handle(self, file_path: Path) -> bool:
        """Check if this handler can handle the file"""
        return file_path.suffix.lower() in self.supported_extensions

    def read_file(self, file_path: Path) -> str:
        """Read document file (extract text)"""
        ext = file_path.suffix.lower()

        if ext == '.pdf':
            return self._read_pdf(file_path)
        elif ext in ['.docx', '.doc']:
            return self._read_docx(file_path)

        return ""

    def write_file(self, file_path: Path, content: str) -> bool:
        """
        Write extracted text back to a document.

        For V1, LightSpeed treats PDF/DOC/DOCX as preview-only. Returning False
        avoids runtime crashes if a caller attempts to save from the universal editor.
        """
        try:
            messagebox.showinfo(
                "Read-Only Document",
                f"{file_path.suffix.upper()} files are preview-only in this editor.\n\n"
                "Export text instead (Toolbar: Export Text) or edit the source in an external tool."
            )
        except Exception:
            pass
        return False

    def create_editor_widget(
        self,
        parent: tk.Widget,
        file_path: Path,
        content: str,
        **kwargs
    ) -> tk.Widget:
        """Create document viewer widget"""
        self.current_file = file_path

        # Split content into pages (for PDF)
        if file_path.suffix.lower() == '.pdf':
            self.current_pages = content.split('\n\n=== PAGE ')
            if len(self.current_pages) > 1:
                self.current_pages = self.current_pages[1:]  # Remove empty first element
        else:
            self.current_pages = [content]

        self.current_page_index = 0

        # Create container
        container = ttk.Frame(parent)

        # Toolbar for page navigation
        if len(self.current_pages) > 1:
            nav_frame = ttk.Frame(container)
            nav_frame.pack(fill=tk.X, pady=5)

            ttk.Button(
                nav_frame,
                text="◀ Previous",
                command=self._prev_page
            ).pack(side=tk.LEFT, padx=2)

            self.page_label = ttk.Label(
                nav_frame,
                text=f"Page 1 of {len(self.current_pages)}"
            )
            self.page_label.pack(side=tk.LEFT, padx=10)

            ttk.Button(
                nav_frame,
                text="Next ▶",
                command=self._next_page
            ).pack(side=tk.LEFT, padx=2)

        # Text widget for content display
        text_frame = ttk.Frame(container)
        text_frame.pack(fill=tk.BOTH, expand=True)

        scrollbar = ttk.Scrollbar(text_frame)
        scrollbar.pack(side=tk.RIGHT, fill=tk.Y)

        text_widget = tk.Text(
            text_frame,
            wrap=tk.WORD,
            yscrollcommand=scrollbar.set,
            font=('Arial', 10),
            state='disabled'  # Read-only
        )
        text_widget.pack(fill=tk.BOTH, expand=True)

        scrollbar.config(command=text_widget.yview)

        # Display first page
        self._display_page(text_widget)

        self.current_widget = text_widget

        # Metadata info
        info_frame = ttk.Frame(container)
        info_frame.pack(fill=tk.X, pady=5)

        metadata = self._get_metadata(file_path)
        info_text = " | ".join([f"{k}: {v}" for k, v in metadata.items()])

        ttk.Label(
            info_frame,
            text=info_text,
            relief=tk.SUNKEN,
            anchor=tk.W
        ).pack(fill=tk.X)

        return container

    def get_content_from_widget(self, widget: tk.Widget) -> str:
        """Get content (not editable, return empty)"""
        return ""

    def get_toolbar_actions(self) -> List[EditorAction]:
        """Get toolbar actions"""
        actions = [
            EditorAction(
                action_id="metadata",
                label="Metadata",
                callback=self._show_metadata,
                tooltip="Show document metadata"
            ),
            EditorAction(
                action_id="export",
                label="Export Text",
                callback=self._export_text,
                tooltip="Export extracted text"
            ),
        ]

        return actions

    def _read_pdf(self, file_path: Path) -> str:
        """Extract text from PDF"""
        try:
            import PyPDF2

            text_pages = []

            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)

                for i, page in enumerate(reader.pages, 1):
                    page_text = page.extract_text()
                    text_pages.append(f"=== PAGE {i} ===\n\n{page_text}")

            return '\n\n'.join(text_pages)

        except ImportError:
            # Try pdfplumber as fallback
            try:
                import pdfplumber

                text_pages = []

                with pdfplumber.open(file_path) as pdf:
                    for i, page in enumerate(pdf.pages, 1):
                        page_text = page.extract_text()
                        text_pages.append(f"=== PAGE {i} ===\n\n{page_text}")

                return '\n\n'.join(text_pages)

            except ImportError:
                messagebox.showwarning(
                    "PDF Support",
                    "PDF reading requires PyPDF2 or pdfplumber.\n" +
                    "Install with: pip install PyPDF2"
                )
                return "[PDF preview requires PyPDF2 or pdfplumber]"

        except Exception as e:
            return f"[Error reading PDF: {e}]"

    def _read_docx(self, file_path: Path) -> str:
        """Extract text from DOCX"""
        try:
            import docx

            doc = docx.Document(file_path)
            paragraphs = [para.text for para in doc.paragraphs]

            return '\n\n'.join(paragraphs)

        except ImportError:
            messagebox.showwarning(
                "DOCX Support",
                "DOCX reading requires python-docx.\n" +
                "Install with: pip install python-docx"
            )
            return "[DOCX preview requires python-docx]"

        except Exception as e:
            return f"[Error reading DOCX: {e}]"

    def _get_metadata(self, file_path: Path) -> dict:
        """Get document metadata"""
        metadata = {}

        ext = file_path.suffix.lower()

        if ext == '.pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    reader = PyPDF2.PdfReader(f)
                    info = reader.metadata

                    if info:
                        if '/Title' in info:
                            metadata['Title'] = str(info['/Title'])
                        if '/Author' in info:
                            metadata['Author'] = str(info['/Author'])

                    metadata['Pages'] = len(reader.pages)

            except:
                pass

        elif ext in ['.docx', '.doc']:
            try:
                import docx
                doc = docx.Document(file_path)
                props = doc.core_properties

                if props.title:
                    metadata['Title'] = props.title
                if props.author:
                    metadata['Author'] = props.author

            except:
                pass

        # File size
        size_bytes = file_path.stat().st_size
        size_kb = size_bytes / 1024

        if size_kb < 1024:
            metadata['Size'] = f"{size_kb:.1f} KB"
        else:
            metadata['Size'] = f"{size_kb/1024:.1f} MB"

        return metadata

    def _display_page(self, text_widget: tk.Text):
        """Display current page"""
        if not self.current_pages:
            return

        text_widget.config(state='normal')
        text_widget.delete('1.0', tk.END)

        page_content = self.current_pages[self.current_page_index]
        text_widget.insert('1.0', page_content)

        text_widget.config(state='disabled')

    def _prev_page(self):
        """Go to previous page"""
        if self.current_page_index > 0:
            self.current_page_index -= 1
            self._display_page(self.current_widget)

            if hasattr(self, 'page_label'):
                self.page_label.config(
                    text=f"Page {self.current_page_index + 1} of {len(self.current_pages)}"
                )

    def _next_page(self):
        """Go to next page"""
        if self.current_page_index < len(self.current_pages) - 1:
            self.current_page_index += 1
            self._display_page(self.current_widget)

            if hasattr(self, 'page_label'):
                self.page_label.config(
                    text=f"Page {self.current_page_index + 1} of {len(self.current_pages)}"
                )

    def _show_metadata(self):
        """Show full metadata"""
        metadata = self._get_metadata(self.current_file)

        metadata_text = "Document Metadata\n\n"
        for key, value in metadata.items():
            metadata_text += f"{key}: {value}\n"

        messagebox.showinfo("Metadata", metadata_text)

    def _export_text(self):
        """Export extracted text"""
        from tkinter import filedialog

        output_path = filedialog.asksaveasfilename(
            title="Export Text",
            defaultextension=".txt",
            filetypes=[("Text Files", "*.txt")]
        )

        if output_path:
            try:
                content = '\n\n'.join(self.current_pages)
                Path(output_path).write_text(content, encoding='utf-8')
                messagebox.showinfo("Success", f"Text exported to {output_path}")
            except Exception as e:
                messagebox.showerror("Error", f"Failed to export: {e}")
